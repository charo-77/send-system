from __future__ import annotations

import base64
import mimetypes
from dataclasses import dataclass
from html import escape
from pathlib import Path
from typing import List
from zipfile import ZipFile
import shutil
import re

from docx import Document
import re


TITLE_PREFIX_RE = re.compile(
    r"^((?:\d{6,}|\d{4}[-_]?\d{2}[-_]?\d{2}[-_]?\d{0,6})(?:__|[_\-\s]+)|(?:[A-Za-z0-9]+__)+|(?:[^_\\/]{1,30}__)+|(?:\d+[._、\-\s]+))+"
)


def normalize_article_title(value: str) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    text = text.replace("　", " ")
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.S).strip()
    text = TITLE_PREFIX_RE.sub("", text)
    parts = text.split("__")
    if len(parts) >= 4 and re.fullmatch(r"\d{8}_\d{6}_\d+", parts[0] or ""):
        text = parts[-1]
    elif len(parts) >= 3 and parts[0].isdigit():
        text = parts[-1]
    text = re.sub(r"^[A-Za-z0-9_-]{1,32}__", "", text)
    text = re.sub(r"^\d+[._、\-\s]+", "", text)
    text = re.sub(r"[\x00-\x1f]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text.strip(" _-，。；：、【】[]()（）《》<>") or str(value or "").strip()


@dataclass
class Article:
    path: Path
    title: str
    body: str
    body_html: str = ""


def _para_has_content(para) -> bool:
    if para.text and para.text.strip():
        return True
    return bool(getattr(para._element, 'xpath', lambda *_: [])('.//*[local-name()="blip"]'))


def _para_to_html(para, rel_to_dataurl: dict[str, str]) -> str:
    parts: list[str] = []
    for run in para.runs:
        text = escape(run.text or "")
        if text:
            if run.bold:
                text = f"<strong>{text}</strong>"
            if run.italic:
                text = f"<em>{text}</em>"
            if run.underline:
                text = f"<u>{text}</u>"
            parts.append(text)
        # inline images in the same run
        blips = getattr(run._element, 'xpath', lambda *_: [])('.//*[local-name()="blip"]')
        for blip in blips:
            rid = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if rid and rid in rel_to_dataurl:
                parts.append(f'<img src="{rel_to_dataurl[rid]}" />')
    return "<p>" + ("".join(parts) if parts else "<br>") + "</p>"


def _docx_image_relmap(doc: Document, path: Path) -> dict[str, str]:
    relmap: dict[str, str] = {}
    with ZipFile(path) as z:
        for rel in doc.part.rels.values():
            try:
                target = str(rel.target_ref)
                if 'image' not in target:
                    continue
                name = target if target.startswith('word/') else f'word/{target}'
                data = z.read(name)
                mime = mimetypes.guess_type(name)[0] or 'image/png'
                relmap[rel.rId] = f"data:{mime};base64," + base64.b64encode(data).decode('ascii')
            except Exception:
                continue
    return relmap


def extract_docx_article(path: Path) -> Article:
    doc = Document(str(path))
    paras = [normalize_article_title(p.text) for p in doc.paragraphs if p.text and p.text.strip()]
    paras = [p for p in paras if p]
    # Prefer the first non-empty paragraph as title; if unavailable, strip processing prefixes from filename.
    title = paras[0] if paras else normalize_article_title(path.stem)
    body_paras = paras[1:] if len(paras) > 1 else []
    body = "\n\n".join(body_paras or paras)
    relmap = _docx_image_relmap(doc, path)
    body_html = "\n".join(_para_to_html(p, relmap) for p in doc.paragraphs if _para_has_content(p))
    return Article(path=path, title=title, body=body, body_html=body_html)


def extract_docx_images(path: Path, out_dir: Path) -> list[Path]:
    """Extract embedded docx images in order for cover selection/upload."""
    out_dir.mkdir(parents=True, exist_ok=True)
    suffix_map = {
        ".jpeg": ".jpg",
        ".jpg": ".jpg",
        ".png": ".png",
        ".webp": ".webp",
    }
    outs: list[Path] = []
    with ZipFile(path) as z:
        media = sorted(n for n in z.namelist() if n.startswith("word/media/"))
        for idx, name in enumerate(media, start=1):
            ext = Path(name).suffix.lower()
            if ext not in suffix_map:
                continue
            out = out_dir / f"{path.stem}_img{idx:02d}{suffix_map[ext]}"
            out.write_bytes(z.read(name))
            outs.append(out)
    return outs


def extract_first_docx_image(path: Path, out_dir: Path) -> Path | None:
    """Extract the first embedded docx image for use as Baijiahao cover."""
    imgs = extract_docx_images(path, out_dir)
    return imgs[0] if imgs else None


def list_docx(root: Path) -> List[Path]:
    skipped_dirs = {"A发布成功", "A发布失败"}
    return sorted([
        p for p in root.rglob("*.docx")
        if p.is_file() and not any(part in skipped_dirs for part in p.parts)
    ])
